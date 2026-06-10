from pathlib import Path

from docx import Document

from app.services.docx_chunk_builder import build_chunks_from_docx
from app.services.structured_table import (
    format_labeled_row,
    structured_table_from_docx_grid,
    table_to_built_chunks,
)


def _make_playbook_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("NDA Review Playbook", level=1)
    table = doc.add_table(rows=3, cols=4)
    headers = ["Sl No.", "Clause", "Preferred positions", "Fallback positions"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "Entity and signatory details"
    table.rows[1].cells[2].text = "Entity Name: MiddleGround Management"
    table.rows[1].cells[3].text = "Same as preferred"
    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "Ownership of Company Interest"
    table.rows[2].cells[2].text = "Delete any representation about voting securities"
    table.rows[2].cells[3].text = "Escalate if pushed"
    doc.save(path)


def test_structured_table_from_docx_grid_uses_literal_headers():
    grid = [
        ["Sl No.", "Clause", "Preferred positions", "Fallback positions"],
        ["1", "Entity and signatory details", "Preferred text", "Fallback text"],
    ]
    table = structured_table_from_docx_grid(
        table_id="t0",
        rows_grid=grid,
        page_number=1,
        heading_path="NDA Review Playbook",
        block_index=1,
    )
    assert table.columns == ["Sl No.", "Clause", "Preferred positions", "Fallback positions"]
    assert table.rows[0]["Clause"] == "Entity and signatory details"
    assert table.parse_confidence >= 0.5


def test_labeled_row_preserves_column_headers():
    row = {"Clause": "Standstill", "Preferred positions": "6 month term"}
    text = format_labeled_row(row, heading_path="Playbook", columns=["Clause", "Preferred positions"])
    assert "Clause: Standstill" in text
    assert "Preferred positions: 6 month term" in text


def test_docx_build_emits_table_row_chunks(tmp_path: Path):
    docx_path = tmp_path / "playbook.docx"
    _make_playbook_docx(docx_path)
    chunks, _pages, meta = build_chunks_from_docx(str(docx_path))
    types = {c.chunk_type for c in chunks}
    assert "table_row" in types
    assert "table_header" in types
    assert meta["table_count"] == 1
    row_chunks = [c for c in chunks if c.chunk_type == "table_row"]
    assert len(row_chunks) >= 2
    assert any("Ownership of Company Interest" in c.text_content for c in row_chunks)
    assert any("Clause:" in c.text_content for c in row_chunks)


def test_table_to_built_chunks_low_confidence_fallback():
    from app.services.structured_table import StructuredTable

    table = StructuredTable(
        table_id="t0",
        source="docx",
        page_start=1,
        page_end=1,
        columns=[],
        header_row_count=0,
        rows=[],
        parse_confidence=0.1,
        heading_path=None,
        block_index=0,
    )
    built = table_to_built_chunks(table, page_number=1, total_blocks=1)
    assert len(built) == 0 or built[0].chunk_type == "table"
