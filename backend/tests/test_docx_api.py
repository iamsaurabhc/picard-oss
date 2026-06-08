import io
import uuid

from docx import Document

from app.db.models import Chunk


def _docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Test Section", level=1)
    doc.add_paragraph("Hello from DOCX upload test.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_upload_docx(client, monkeypatch):
    monkeypatch.setattr("app.services.ingestion._executor.submit", lambda *args, **kwargs: None)
    ws = client.post("/workspaces", json={"name": "DOCX Upload"}).json()
    files = {
        "file": (
            "memo.docx",
            _docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    r = client.post(f"/workspaces/{ws['id']}/documents", files=files)
    assert r.status_code == 200
    body = r.json()
    assert body["file_type"] == "docx"
    assert body["file_name"] == "memo.docx"


def test_reject_legacy_doc(client):
    ws = client.post("/workspaces", json={"name": "Reject"}).json()
    files = {"file": ("old.doc", b"fake", "application/msword")}
    r = client.post(f"/workspaces/{ws['id']}/documents", files=files)
    assert r.status_code == 400


def test_convert_pdf_to_docx_chunks(client, monkeypatch, db_session):
    from app.db.models import Chunk, Document as DbDocument

    monkeypatch.setattr("app.services.ingestion._executor.submit", lambda *args, **kwargs: None)
    ws = client.post("/workspaces", json={"name": "Convert"}).json()
    pdf_bytes = b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF"
    files = {"file": ("scan.pdf", pdf_bytes, "application/pdf")}
    doc = client.post(f"/workspaces/{ws['id']}/documents", files=files).json()

    # Simulate parsed chunks for chunk-based conversion

    row = db_session.get(DbDocument, doc["id"])
    row.parse_status = "done"
    row.text_source = "scanned"
    db_session.add(
        Chunk(
            id=str(uuid.uuid4()),
            document_id=doc["id"],
            page_number=1,
            chunk_type="paragraph",
            bbox_json='{"x0":0,"y0":0,"x1":1,"y1":1}',
            text_content="Scanned line one.",
            heading_path=None,
            section_key=None,
            token_count=3,
        )
    )
    db_session.commit()

    r = client.post(f"/documents/{doc['id']}/convert-to-docx?method=chunks")
    assert r.status_code == 200
    body = r.json()
    assert body["method"] == "chunks"
    assert body["file_name"].endswith(".docx")


def test_update_docx_file(client, monkeypatch, db_session):
    monkeypatch.setattr("app.services.ingestion._executor.submit", lambda *args, **kwargs: None)
    ws = client.post("/workspaces", json={"name": "Save"}).json()
    files = {
        "file": (
            "edit.docx",
            _docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    doc = client.post(f"/workspaces/{ws['id']}/documents", files=files).json()

    updated = Document()
    updated.add_paragraph("Updated paragraph content for re-index.")
    buf = io.BytesIO()
    updated.save(buf)
    new_bytes = buf.getvalue()

    files2 = {
        "file": (
            "edit.docx",
            new_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    r = client.put(f"/documents/{doc['id']}/file", files=files2)
    assert r.status_code == 200
    assert r.json()["parse_status"] == "pending"
    assert r.json()["content_hash"] != doc["content_hash"]
