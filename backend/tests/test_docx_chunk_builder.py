import io
import uuid
from pathlib import Path

from docx import Document

from app.db.models import Chunk, Document as DbDocument, Workspace
from app.db.session import utc_now_iso
from app.services.docx_chunk_builder import build_chunks_from_docx
from app.services.pdf_to_docx import chunks_to_docx_bytes


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Condition A", level=1)
    doc.add_paragraph("Party ABC obligations under the agreement.")
    doc.add_paragraph("- First list item")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Col A"
    table.rows[0].cells[1].text = "Col B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    doc.save(path)


def test_build_chunks_from_docx(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    _make_docx(docx_path)

    chunks, page_count, meta = build_chunks_from_docx(str(docx_path))

    assert page_count >= 1
    assert meta["text_source"] == "docx"
    assert len(chunks) >= 3
    types = {c.chunk_type for c in chunks}
    assert "heading" in types
    assert "paragraph" in types or "list" in types
    assert any(c.heading_path and "Condition A" in c.heading_path for c in chunks)
    assert all(c.anchor_json for c in chunks)


def test_chunks_to_docx_bytes_roundtrip(db_session, tmp_path):
    now = utc_now_iso()
    ws = Workspace(id=str(uuid.uuid4()), name="DOCX", matter_ref=None, created_at=now, updated_at=now)
    db_session.add(ws)
    doc_id = str(uuid.uuid4())
    db_session.add(
        DbDocument(
            id=doc_id,
            workspace_id=ws.id,
            file_name="scan.pdf",
            local_path="pdfs/x/y.pdf",
            parse_status="done",
            created_at=now,
        )
    )
    db_session.add(
        Chunk(
            id=str(uuid.uuid4()),
            document_id=doc_id,
            page_number=1,
            chunk_type="heading",
            bbox_json='{"x0":0,"y0":0,"x1":1,"y1":0.1}',
            text_content="Condition A",
            heading_path="Condition A",
            section_key="abc",
            token_count=2,
        )
    )
    db_session.add(
        Chunk(
            id=str(uuid.uuid4()),
            document_id=doc_id,
            page_number=1,
            chunk_type="paragraph",
            bbox_json='{"x0":0,"y0":0.2,"x1":1,"y1":0.3}',
            text_content="OCR paragraph text.",
            heading_path="Condition A",
            section_key="abc",
            token_count=3,
        )
    )
    db_session.commit()

    out_path = tmp_path / "from_chunks.docx"
    out_path.write_bytes(
        chunks_to_docx_bytes(db_session.query(Chunk).filter(Chunk.document_id == doc_id).all())
    )

    chunks, _, _ = build_chunks_from_docx(str(out_path))
    texts = " ".join(c.text_content for c in chunks)
    assert "Condition A" in texts
    assert "OCR paragraph" in texts
