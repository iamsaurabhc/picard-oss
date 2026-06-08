from __future__ import annotations

import io
import json
import re
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_BREAK
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.db.models import Chunk

if TYPE_CHECKING:
    pass

_LIST_RE = re.compile(r"^[\-\*•]\s+")


def _sort_chunks(chunks: list[Chunk]) -> list[Chunk]:
    def sort_key(c: Chunk) -> tuple[int, float, float]:
        try:
            bbox = json.loads(c.bbox_json)
            y0 = float(bbox.get("y0", 0))
            x0 = float(bbox.get("x0", 0))
        except (json.JSONDecodeError, TypeError, ValueError):
            y0, x0 = 0.0, 0.0
        return (c.page_number, y0, x0)

    return sorted(chunks, key=sort_key)


def _heading_level(heading_path: str | None) -> int:
    if not heading_path:
        return 1
    return min(9, max(1, len(heading_path.split(" > "))))


def _add_list_paragraph(doc: Document, text: str) -> None:
    body = text.strip()
    if _LIST_RE.match(body):
        body = _LIST_RE.sub("", body, count=1).strip()
    doc.add_paragraph(body, style="List Bullet")


def chunks_to_docx_bytes(chunks: list[Chunk]) -> bytes:
    """Assemble an editable DOCX from liteparse chunks (scanned/mixed PDFs)."""
    doc = Document()
    sorted_chunks = _sort_chunks(chunks)
    last_page: int | None = None

    for chunk in sorted_chunks:
        if last_page is not None and chunk.page_number != last_page:
            para = doc.add_paragraph()
            para.add_run().add_break(WD_BREAK.PAGE)
        last_page = chunk.page_number

        text = chunk.text_content.strip()
        if not text:
            continue

        if chunk.chunk_type == "heading":
            level = _heading_level(chunk.heading_path)
            doc.add_heading(text, level=level)
        elif chunk.chunk_type == "list":
            _add_list_paragraph(doc, text)
        elif chunk.chunk_type == "table":
            rows = [line.split("\t") for line in text.split("\n") if line.strip()]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < len(table.rows[r_idx].cells):
                            table.rows[r_idx].cells[c_idx].text = cell_text
            else:
                doc.add_paragraph(text)
        else:
            doc.add_paragraph(text)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def chunks_to_docx_for_document(db: Session, document_id: str) -> bytes:
    chunks = list(db.scalars(select(Chunk).where(Chunk.document_id == document_id)).all())
    if not chunks:
        raise ValueError("Document has no chunks — parse the PDF first")
    return chunks_to_docx_bytes(chunks)
